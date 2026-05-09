from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}


@dataclass(frozen=True)
class ParsedSection:
    text: str
    page: int | None = None
    label: str | None = None


def parse_document(path: Path) -> list[ParsedSection]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in {".md", ".markdown", ".txt"}:
        return _parse_text(path)
    if ext == ".docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")


def _parse_text(path: Path) -> list[ParsedSection]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [ParsedSection(text=_normalize_text(text), label=path.name)]


def _parse_pdf(path: Path) -> list[ParsedSection]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf.") from exc

    reader = PdfReader(str(path))
    sections: list[ParsedSection] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = _normalize_text(text)
        if text:
            sections.append(ParsedSection(text=text, page=index, label=f"page {index}"))
    return sections


def _parse_docx(path: Path) -> list[ParsedSection]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx.") from exc

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return [ParsedSection(text=_normalize_text("\n".join(lines)), label=path.name)]


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    normalized: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank:
                normalized.append("")
            blank = True
            continue
        normalized.append(line)
        blank = False
    return "\n".join(normalized).strip()

